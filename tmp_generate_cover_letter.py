from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


src = Path("/Users/fangyuanyuan/Downloads/论文/cover letter.docx")
out = Path(
    "/Users/fangyuanyuan/Downloads/EMR数据缺省综述/补全方法论文/"
    "ClinDiff-PDAC_cover_letter.docx"
)

doc = Document(str(src))

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.text = ""

texts = [
    "Dear Editor,",
    (
        "We are pleased to submit our manuscript entitled “Targeted imputation of "
        "bilirubin-linked liver biomarkers in pancreatic cancer electronic health "
        "records: a single-centre proof-of-concept study” for consideration as an "
        "Article in Communications Medicine."
    ),
    "Innovation Highlights",
    "Innovation 1. A focused clinical missing-data problem in pancreatic cancer EHRs",
    (
        "We studied 2,347 de-identified pancreatic ductal adenocarcinoma electronic "
        "health records from Shanghai General Hospital collected between 2018 and "
        "2023. Rather than proposing a broad imputation platform, the manuscript "
        "addresses a narrow but clinically meaningful problem: targeted imputation "
        "for the physiologically linked liver-biomarker trio of total bilirubin, "
        "direct bilirubin and gamma-glutamyl transferase."
    ),
    "Innovation 2. A lightweight, missingness-informed ClinDiff-Gated workflow",
    (
        "ClinDiff-Gated combines four-state missingness representation, feature-wise "
        "conservative routing, conditional liver-trio refinement and simple clinical "
        "constraints such as non-negativity and DB <= TB. The headline workflow is "
        "intentionally lightweight and auditable, and the manuscript distinguishes "
        "the implemented primary workflow from optional repository modules that were "
        "not used for headline claims."
    ),
    "Innovation 3. Repeated-masking evidence with conservative external-validity boundaries",
    (
        "The primary evaluation used prespecified 100-seed MCAR repeated masking at "
        "20%, 40% and 60% missingness, with liver-trio RMSE as the primary endpoint "
        "and non-outcome overall RMSE as a secondary endpoint. The results showed "
        "modest, targeted RMSE reductions concentrated in bilirubin-linked biomarkers "
        "under recoverable missingness, while the manuscript explicitly treats SEER "
        "and TCGA-PAAD as external contextualization rather than external "
        "imputation-accuracy validation because they lack compatible TB, DB and GGT "
        "laboratory panels."
    ),
    (
        "We believe this study is well aligned with the translational scope of "
        "Communications Medicine because it addresses a practical data-quality "
        "barrier in oncology EHR research while maintaining a cautious interpretation "
        "of simulation-based evidence. The work is intended as a single-centre "
        "proof-of-concept study and identifies the next validation step: same-workflow "
        "replay in independent cohorts with compatible liver laboratory panels."
    ),
    (
        "The use of the de-identified Shanghai General Hospital EHR cohort was "
        "approved by the Institutional Review Board of Shanghai General Hospital "
        "(approval number: 院伦快【2025】615号), conducted in accordance with the "
        "Declaration of Helsinki, and granted a waiver of informed consent for "
        "retrospective record review. The code package is publicly available at "
        "https://github.com/anitafangyy-create/clindiff-pdac; patient-level EHR data "
        "are not publicly deposited because they contain sensitive oncology records "
        "and remain subject to institutional data-governance review."
    ),
    (
        "We confirm that this manuscript is original, has not been published "
        "elsewhere, and is not under consideration by another journal. All authors "
        "have approved the submission. Thank you for your consideration. We look "
        "forward to your response."
    ),
    "Prof. Kaishun Wu",
    "The Hong Kong University of Science and Technology (Guangzhou), wuks@hkust-gz.edu.cn",
    "Prof. Qi Li",
    (
        "Cancer Center, Shanghai General Hospital, Shanghai Jiao Tong University "
        "School of Medicine; Shanghai Key Laboratory of Pancreatic Disease, Institute "
        "of Pancreatic Disease, Shanghai Jiao Tong University School of Medicine, "
        "leeqi@sjtu.edu.cn"
    ),
    "Prof. Irfan Ahmed",
    "King's College Hospital London - Saudi Arabia, Irfan.Ahmed@kch.sa",
]

while len(doc.paragraphs) < len(texts):
    doc.add_paragraph("")

for idx, text in enumerate(texts):
    paragraph = doc.paragraphs[idx]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if text == "Innovation Highlights" or text.startswith("Innovation "):
        run.bold = True
    if text.startswith("Prof. "):
        run.bold = True

for paragraph in doc.paragraphs[len(texts) :]:
    for run in paragraph.runs:
        run.text = ""

for section in doc.sections:
    section.top_margin = 914400
    section.bottom_margin = 914400
    section.left_margin = 914400
    section.right_margin = 914400
    for paragraph in section.footer.paragraphs:
        for run in paragraph.runs:
            run.text = ""

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)

doc.save(str(out))
print(out)
